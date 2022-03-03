# wordを扱うモジュール
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT


class Wordopration():
    # ワードに記載する初期設定
    def __init__(self):
        self.new_document = Document()

    def add_title(self,text,size,under=False):
        # 見出しのスタイル設定
        add_paragraph = self.new_document.add_paragraph()
        paragraph_format = add_paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addtitle = add_paragraph.add_run(text)
        addtitle.font.size = Pt(size)
        addtitle.underline = under
        add_paragraph.add_run().add_break()

    def add_date(self,text,size=11,under=False):
        # 日付のスタイル設定
        add_paragraph = self.new_document.add_paragraph()
        paragraph_format = add_paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        adddate = add_paragraph.add_run(text)
        adddate.font.size = Pt(size)
        adddate.underline = under
        add_paragraph.add_run().add_break()

    def add_table(self,group,member,member_num):
        # 表紙の初期設定
        table = self.new_document.add_table(rows=5,cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0,1).text = '報告者'
        table.cell(0,3).text = '学籍番号'
        table.cell(0,4).text = '氏名'
        table.cell(0,1).text = ''
        table.cell(1,2).text = '班'
        table.cell(1,1).text = group
        
        for num in range(len(member)):
            # 学籍番号、氏名を書き込み
            table.cell(num+1,3).text = member_num[num].get()
            table.cell(num+1,4).text = member[num].get()
        
    def add_text(self,text,size=11,under=False,kaigyou=False):
        # pdfから読み取ったテキストをワードに書き込み
        add_paragraph = self.new_document.add_paragraph()
        addtext = add_paragraph.add_run(text)
        addtext.font.size = Pt(size)
        addtext.underline = under
        if kaigyou==True:
            add_paragraph.add_run().add_break()
